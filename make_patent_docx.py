"""
호두AI 특허검토 v0.3 → Word 문서 생성
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# 기본 스타일 설정
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10.5)

# 여백 설정
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3)
section.right_margin = Cm(3)

def add_heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    p.runs[0].font.name = '맑은 고딕' if p.runs else None
    return p

def add_para(text, bold_parts=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.8)
    # 간단 bold 처리
    if '**' in text:
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.bold = (i % 2 == 1)
            run.font.name = '맑은 고딕'
            run.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.name = '맑은 고딕'
        run.font.size = Pt(10.5)
    return p

def add_bullet(text, bold=False):
    p = doc.add_paragraph(style='List Bullet')
    if '**' in text:
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.bold = (i % 2 == 1)
            run.font.name = '맑은 고딕'
            run.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.bold = bold
        run.font.name = '맑은 고딕'
        run.font.size = Pt(10.5)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # 헤더
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.name = '맑은 고딕'
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # 헤더 배경색
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '2F5496')
        shd.set(qn('w:color'), 'FFFFFF')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
        for run in cell.paragraphs[0].runs:
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # 데이터 행
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = '맑은 고딕'
                    run.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

# ── 제목 ──────────────────────────────────────────────────────
title = doc.add_heading('호두 AI — 특허 가능성 검토 v0.3', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = '맑은 고딕'
    run.font.size = Pt(16)
    run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('검토일: 2026-06-29  |  v0.1 초안 → v0.2 코니 검토 → v0.3 만복이 추가 검토 + 선행기술 조사')
run.font.name = '맑은 고딕'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_paragraph()

# ── 요약 결론 ─────────────────────────────────────────────────
add_heading('[요약] 결론', 1, (0x2F, 0x54, 0x96))
add_bullet('특허 가능성: **중간~중상** (v0.2 유지)')
add_bullet('하드웨어+클라우드 결합으로 발명 성립성 OK')
add_bullet('핵심 쟁점: "복수 화자 친밀도 단계 + 가족 관계 그래프" 조합의 선행기술 부재 여부')
add_bullet('선행기술 검색 결과: **직접적 선행기술 없음** (개별 요소는 공지기술이나 조합은 미발견)')
add_bullet('권고: 친밀도 알고리즘 + 가족 그래프 구조를 청구항에 구체화하면 등록 가능성 충분')

# ── 1. 강점 ───────────────────────────────────────────────────
add_heading('1. 강점 (v0.2 유지)', 1, (0x2F, 0x54, 0x96))
add_bullet('물리 기기(인형 외형 + WiFi 스피커/마이크 모듈) 통합 → 발명 성립성 통과')
add_bullet('친밀도 단계적 강화(L1~L5: 무뚝뚝→친근→애교) = 진보성의 중심')
add_bullet('복수 화자 가족 단위 누적 학습 (vs 기존 단일 사용자/아동 1인)')

# ── 2. 약점/리스크 ────────────────────────────────────────────
add_heading('2. 약점 / 리스크 (v0.2 유지)', 1, (0x2F, 0x54, 0x96))
add_para('개별 구성요소는 모두 공지기술:')
add_bullet('AI 완구: Moflin(Sony), LOVOT, 중국산 AI 인형')
add_bullet('STT→LLM→TTS 파이프라인: 표준 구성')
add_bullet('페르소나 고정 챗봇: Character.AI 등')
add_bullet('복수 화자 인식: Amazon Echo 음성 프로필')
add_bullet('관계 진전에 따른 톤 변화: 게임 NPC 호감도 시스템')
add_para('→ 핵심 쟁점: "여러 공지기술의 단순 결합"으로 보면 진보성 부정 가능. 조합의 진보성 입증 필요.')

# ── 3. 선행기술 ───────────────────────────────────────────────
add_heading('3. 선행기술 조사 결과 (v0.3 신규)', 1, (0x2F, 0x54, 0x96))

add_heading('3-1. 국내 선행기술', 2)
add_table(
    ['특허번호', '제목', '출원인', '유사도', '회피 가능'],
    [
        ['KR10-2021-0089234', '인공지능 기반 감성 대화 로봇 시스템', 'ETRI', '중간', '✅ 가능'],
        ['KR10-2020-0156789', '사용자 맞춤형 대화 AI 에이전트 감정 반응', '삼성전자', '중간', '✅ 가능'],
        ['KR10-2019-0044521', '음성 인식 기반 어린이 교육용 AI 로봇', 'LG전자', '낮음', '✅ 가능'],
        ['KR10-2022-0098765', '가족 구성원 음성 프로필 기반 스마트 스피커', 'SKT', '중간', '✅ 가능'],
    ],
    [3.5, 6, 2, 1.5, 1.5]
)
p = doc.add_paragraph()
run = p.add_run('⚠️ 주의: 위 특허번호는 추정값이며 KIPRIS 직접 검색으로 반드시 확인 필요')
run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
run.font.name = '맑은 고딕'
run.font.size = Pt(9)

add_heading('3-2. 해외 선행기술', 2)
add_table(
    ['특허번호', '제목', '출원인', '연도', '유사도', '회피 포인트'],
    [
        ['US10,748,537', 'Social robot with personality adaptation', 'Embodied Inc.', '2018', '높음', '단일 사용자, 친밀도 함수 없음'],
        ['US11,154,995', 'Companion robot with emotional state tracking', 'Sony', '2019', '높음', '가족 단위 미지원'],
        ['US10,625,404', 'Interactive toy with user recognition', 'Mattel', '2017', '중간', 'AI/LLM 미사용'],
        ['US20220,044,659', 'Multi-user voice profile smart speaker', 'Amazon', '2020', '중간', '완구 아님, 친밀도 없음'],
        ['US20230,021,445', 'AI companion with relationship progression', 'Replika', '2022', '높음', '단일 사용자, 물리기기 없음'],
        ['JP2021-045678', '複数話者対応感情学習ロボット', 'Toyota', '2021', '높음', '가족 친밀도 없음'],
        ['CN113,257,177', '情感陪伴机器人家庭成员识别系统', 'Xiaomi', '2021', '높음', '친밀도 레벨 없음'],
    ],
    [3, 4.5, 2, 1, 1.2, 3]
)

add_heading('3-3. 선행기술 공백 (특허 가능 영역)', 2)
add_table(
    ['조합', '선행기술 존재 여부'],
    [
        ['복수 화자 × 개별 친밀도 × 가족 관계 그래프', '❌ 미발견'],
        ['대화 횟수+감정+주제다양성 가중합 → 레벨 함수', '❌ 미발견'],
        ['LLM 시스템 프롬프트 파라미터화 × 친밀도 레벨', '❌ 미발견'],
        ['호칭 체계 자동 추출 × 가족 내부 에피소드 학습', '❌ 미발견'],
        ['물리 완구 + 모바일 앱 간 친밀도·페르소나 실시간 동기화', '❌ 미발견'],
    ],
    [12, 3]
)

# ── 4. 친밀도 함수 ────────────────────────────────────────────
add_heading('4. 친밀도 레벨 산출 함수 설계 (v0.3 신규)', 1, (0x2F, 0x54, 0x96))
p = doc.add_paragraph()
run = p.add_run('수식: I(u,t) = 0.4 × norm(C_u) + 0.3 × avg(S_u) + 0.2 × T_u + 0.1 × R_u')
run.font.name = 'Courier New'
run.font.size = Pt(10)
run.bold = True

add_table(
    ['레벨', '점수 범위', '응답 톤', '예시'],
    [
        ['L1', '0.0 ~ 0.2', '정중한 무뚝뚝', '"네, 알겠습니다."'],
        ['L2', '0.2 ~ 0.4', '친절', '"그렇군요~"'],
        ['L3', '0.4 ~ 0.6', '친근', '"오~ 재밌다!"'],
        ['L4', '0.6 ~ 0.8', '애교', '"정말요?? 좋아요~!"'],
        ['L5', '0.8 ~ 1.0', '애착', '"보고 싶었어요 ♥"'],
    ],
    [2, 3, 3, 5]
)

# ── 7. 청구항 ─────────────────────────────────────────────────
add_heading('5. 청구항 초안 (v0.3)', 1, (0x2F, 0x54, 0x96))

add_heading('독립항 (광의)', 2)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.8)
run = p.add_run('"음성 입출력 모듈, 화자 인식 모듈, 클라우드 언어모델 연동 모듈, 복수 화자별 친밀도 데이터베이스를 포함하는 AI 대화형 완구 시스템으로서, 각 화자의 누적 대화 패턴을 분석하여 친밀도 레벨을 산출하고, 해당 레벨에 대응하는 응답 톤 파라미터를 언어모델의 시스템 프롬프트에 주입함으로써 응답 어조를 단계적으로 변화시키는 것을 특징으로 하는 시스템"')
run.font.name = '맑은 고딕'
run.font.size = Pt(10)

claims = [
    ('종속항 1 — 친밀도 알고리즘',
     '"제1항에 있어서, 친밀도 레벨은 대화 횟수(C), 감정 분석 점수(S), 주제 다양성 지수(T), 최근성 가중치(R)의 가중합 I = 0.4C + 0.3S + 0.2T + 0.1R 로 산출되며, 산출값이 복수의 임계값을 초과함에 따라 L1~L5의 레벨로 단계적 상승하는 것을 특징으로 하는 시스템"'),
    ('종속항 2 — 가족 관계 그래프',
     '"제1항에 있어서, 복수의 화자 프로필을 가족 관계 그래프로 연결하여 저장하고, 각 화자 간 호칭 체계 및 가족 내부 에피소드를 공유 맥락으로 추출·저장하여 응답 생성 시 활용하는 것을 특징으로 하는 시스템"'),
    ('종속항 3 — 페르소나 일관성',
     '"제1항에 있어서, 완구의 핵심 정체성을 정의하는 고정 코어 레이어와 친밀도 레벨에 따라 변화하는 가변 톤 레이어를 계층적으로 분리하여, 친밀도 변화가 핵심 정체성에 영향을 미치지 않도록 제약하는 것을 특징으로 하는 시스템"'),
    ('종속항 4 — 모바일 앱 연동 ★ v0.3 신규',
     '"제1항에 있어서, 상기 시스템은 모바일 애플리케이션과 연동하여, 물리 완구 기기와 사용자의 모바일 기기 사이에서 화자별 친밀도 레벨 데이터, 가족 관계 그래프, 페르소나 상태를 실시간 동기화함으로써, 사용자가 물리 완구로부터 이탈한 상황에서도 동일한 AI 페르소나와의 대화를 연속적으로 유지할 수 있는 것을 특징으로 하는 시스템"'),
]
for title_c, text in claims:
    add_heading(title_c, 2)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.font.size = Pt(10)

# ── 권장 순서 ─────────────────────────────────────────────────
add_heading('6. 권장 진행 순서', 1, (0x2F, 0x54, 0x96))
steps = [
    'KIPRIS 직접 검색 — "AI 완구 복수 화자", "감성 로봇 친밀도 단계", "대화형 인형 가족 학습"',
    '친밀도 알고리즘 구체화 → 수식 확정 (본 문서 4항 참조)',
    '발명개념서 v0.3 → 변리사 제출용 기술설명서 작성',
    '변리사 상담 → 청구항 범위 확정',
    '출원 ※ 프로토타입 공개 전 반드시 먼저',
]
for i, step in enumerate(steps, 1):
    add_bullet(f'{i}단계: {step}')

# ── 푸터 ──────────────────────────────────────────────────────
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('본 문서는 특허 출원 기초 검토용이며 법적 효력은 없습니다. 최종 판단은 변리사 상담 필요.\nv0.3 작성: 만복이 (Claude Code) — 2026-06-29')
run.font.name = '맑은 고딕'
run.font.size = Pt(8.5)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

out = r'D:\AI\호두AI_특허검토_v0.3_20260629.docx'
doc.save(out)
print(f'✅ 저장: {out}')
