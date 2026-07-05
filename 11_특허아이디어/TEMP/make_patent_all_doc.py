# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading('바로보기+만복이 특허 후보 개념서 모음', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = doc.add_paragraph()
meta.add_run('작성일: 2026-06-29  |  작성자: 이한복 (barobogi)  |  버전: v0.1').italic = True
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

items = [
    {
        'title': '① AI 세션 간 자동 컨텍스트 동기화 시스템',
        'prior': 'LangChain Memory, Claude Projects 등 단일 플랫폼 내 메모리. 이기종 AI 간 파일 기반 자동 전달은 없음.',
        'diff': '파일 감시(watchdog) 기반 자동 트리거 + 복수 이기종 AI 에이전트 간 표준화 컨텍스트 전달 + 플랫폼 독립',
        'possibility': '중간 — 순수 SW라 어려우나 시스템+방법 결합 청구항으로 가능성 있음',
        'claim': '파일 시스템 감시 모듈이 AI 세션 종료를 감지하여 다음 AI 에이전트 세션이 판독 가능한 구조화된 컨텍스트 파일을 자동 생성하는 방법 및 시스템',
    },
    {
        'title': '② 외부 AI 위임 실행 패턴 (Multi-tier AI Delegation)',
        'prior': 'AutoGPT, LangChain Agents 등 단일 모델 내 툴 사용. 비용 등급별 AI 역할 분리 개념 없음.',
        'diff': '비용 등급별 AI 역할 명시적 분리(고비용=판단, 저비용=반복) + 파일 인터페이스로 플랫폼 독립 + 2단계 품질 보증',
        'possibility': '중간 — 비용 기반 AI 역할 분리+파일 인터페이스 조합은 신규성 주장 가능',
        'claim': '제1 AI 모델이 작업 명세를 생성하고, 비용이 낮은 제2 AI 모델이 반복 실행 후 결과 파일을 저장하며, 제1 AI 모델이 결과를 검증하는 다단계 AI 실행 위임 방법',
    },
    {
        'title': '③ 시장 결과 기반 투자 신호 자동 학습 시스템',
        'prior': 'RL 기반 트레이딩(대규모 인프라 필요), 알고리즘 백테스팅(사후 검증), Bayesian 최적화(전문가 지식 필요)',
        'diff': '인간 판단 완전 배제 + 시장 가격만이 유일한 학습 신호 + 개인 투자자 수준 경량 파이프라인 + 단순 명확한 피드백 기준(5일 수익률)',
        'possibility': '중간~높음 — 금융+SW 결합 BM특허 가능. 인간 편향 배제라는 기술 효과가 신규성에 유리',
        'claim': '투자 신호 생성 후 일정 기간 경과 시 시장 가격 데이터를 자동 수집하여 인간 판단 개입 없이 신호 패턴별 가중치를 갱신하는 자동화된 투자 신호 최적화 방법',
    },
    {
        'title': '④ 인간-AI 협업을 위한 4계층 지속성 컨텍스트 아키텍처',
        'prior': 'CLAUDE.md, .cursorrules 등 단일 레이어 규칙 파일. RAG는 외부 지식 검색이며 협업 구조 아님.',
        'diff': '4계층 명시적 역할 분리(전역규칙/기술이력/재사용코드/프로젝트현황) + 계층 간 승격 메커니즘 + AI 세션 단절 문제의 시스템적 해결',
        'possibility': '낮음~중간 — 추상적 아키텍처에 가까워 방어적 공개(defensive publication) 적합',
        'claim': '전역 규칙, 기술 이력, 재사용 코드, 프로젝트 현황의 4계층으로 구성된 파일 시스템 기반 인간-AI 협업 컨텍스트 관리 방법 및 시스템',
    },
]

for item in items:
    doc.add_heading(item['title'], 1)

    doc.add_paragraph().add_run('선행 기술 동향').bold = True
    doc.add_paragraph(item['prior'])

    doc.add_paragraph().add_run('차별점').bold = True
    doc.add_paragraph(item['diff'])

    doc.add_paragraph().add_run('특허 가능성').bold = True
    doc.add_paragraph(item['possibility'])

    doc.add_paragraph().add_run('핵심 청구항 후보').bold = True
    doc.add_paragraph(f'"{item["claim"]}"')
    doc.add_paragraph()

doc.add_heading('종합 우선순위', 1)
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '순위'
hdr[1].text = '아이템'
hdr[2].text = '추천 액션'
rows_data = [
    ('1위', '73 호두 AI (HW+SW 결합)', '변리사 상담 → KR 출원 최우선'),
    ('2위', '③ 시장 기반 자동학습', 'KIPRIS 검색 후 검토'),
    ('3위', '② 외부 AI 위임 패턴', 'KIPRIS 검색 후 검토'),
    ('4위', '① AI 세션 동기화', 'KIPRIS 검색 후 검토'),
    ('5위', '④ 4계층 아키텍처', '방어적 공개 검토'),
]
for i, (a, b, c) in enumerate(rows_data, 1):
    r = table.rows[i].cells
    r[0].text = a; r[1].text = b; r[2].text = c

doc.add_paragraph()
p = doc.add_paragraph('* 본 문서는 특허 출원을 위한 기초 개념 정리용이며, 법적 효력은 없습니다.')
p.italic = True

doc.save('D:/AI/DEV_ITEM/특허후보_개념서_모음_v0.1.docx')
print('완료')
