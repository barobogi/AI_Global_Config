# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

t = doc.add_heading('11_#7 — 퍼스널 멀티채널 AI 오케스트레이션 시스템', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = doc.add_paragraph()
meta.add_run('작성일: 2026-06-29  |  작성자: 이한복 (barobogi)  |  버전: v0.21 (요약본)').italic = True
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

doc.add_heading('발명명', 2)
doc.add_paragraph().add_run('모바일 단말기 기반 퍼스널 멀티채널 AI 오케스트레이션 시스템').bold = True

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('한 줄 요약: ').bold = True
p.add_run('"스마트폰을 중심으로 텔레그램·데스크탑 등 이기종 AI 채널이 하나의 AI로 작동하는 시스템"')

doc.add_heading('핵심 차별점 (선행기술 미발견)', 2)
for d in [
    '① 이기종 벤더 AI 채널 간 단일 페르소나 유지 (Amazon/Google은 자사 에코시스템만)',
    '② AI가 상황 분석 후 최적 채널 자동 라우팅 (기존: 사용자가 직접 선택)',
    '③ 파일 시스템 기반 플랫폼 독립 컨텍스트 동기화 (Apple Handoff는 동일 벤더만)',
]:
    doc.add_paragraph(d)

doc.add_heading('HW+SW 결합 구조', 2)
doc.add_paragraph(
    '스마트폰(HW) + 모바일 앱(SW)  ←  HW 결합으로 발명 성립성 확보\n'
    '  ↕\n'
    '텔레그램 봇 (알림/모니터링)\n'
    '  ↕\n'
    '데스크탑 AI CLI (깊은 작업)\n'
    '  ↕\n'
    '[공유 컨텍스트 + AI 자동 라우팅]'
)

doc.add_heading('핵심 청구항', 2)
doc.add_paragraph(
    '"모바일 단말기에 설치된 AI 에이전트 앱, 메신저 봇, 데스크탑 AI를 포함하는 복수의 이기종 AI 채널이 '
    '공유 컨텍스트를 통해 단일 AI 페르소나를 유지하며, AI 라우팅 모듈이 상황에 따라 최적 채널을 자동 선택하고 '
    '채널 전환 시 컨텍스트가 자동 인계되는 퍼스널 AI 오케스트레이션 시스템"'
)

doc.add_heading('11_#1 호두 AI와의 관계', 2)
doc.add_paragraph('• 11_#1 호두 AI: HW = 캐릭터 인형+스피커 → 별개 독립 특허')
doc.add_paragraph('• 11_#7 본 발명: HW = 스마트폰 단말기 → 별개 독립 특허')
doc.add_paragraph('• 향후 통합: 호두 AI + 멀티채널 → 종속항으로 연결 가능')

doc.add_heading('특허 가능성', 2)
doc.add_paragraph('중간~높음 — 스마트폰 HW 결합으로 발명 성립. 이기종 채널 통합+자동라우팅 조합 선행기술 미발견.')

doc.add_heading('다음 단계', 2)
for s in ['① KIPRIS 선행검색', '② 변리사 상담 (11_#1과 별개 출원 전략)', '③ 국내 출원 (KR)']:
    doc.add_paragraph(f'  {s}')

doc.add_paragraph()
doc.add_paragraph('* 상세 내용: 11_7_멀티채널AI_오케스트레이션_v0.2.docx 참조  |  법적 효력 없음').italic = True

doc.save('D:/AI/11_특허아이디어/11_7_멀티채널AI_오케스트레이션_v0.21_요약.docx')
print('v0.21 완료')
