# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

t = doc.add_heading('11_#7 — 퍼스널 멀티채널 AI 오케스트레이션 시스템', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = doc.add_paragraph()
meta.add_run('작성일: 2026-06-29  |  작성자: 이한복 (barobogi)  |  버전: v0.2').italic = True
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('버전 이력: ').bold = True
p.add_run('v0.1 아이디어 초안 → v0.2 스마트폰 HW 결합 확인 + 독립 특허 가능성 확정')
doc.add_paragraph()

doc.add_heading('1. 발명의 명칭', 1)
doc.add_paragraph().add_run('모바일 단말기 기반 퍼스널 멀티채널 AI 오케스트레이션 시스템').bold = True
doc.add_paragraph('(영문: Mobile Device-based Personal Multi-channel AI Orchestration System)')

doc.add_heading('2. 기술 분야', 1)
doc.add_paragraph(
    '스마트폰 단말기(모바일 앱)를 포함한 복수의 이기종 AI 에이전트 채널이 '
    '단일 사용자 컨텍스트와 AI 페르소나를 공유하며, 채널별 역할 분리 및 '
    'AI 자동 라우팅을 통해 기기 전환 시에도 사용자 경험의 연속성을 보장하는 '
    'AI 오케스트레이션 시스템에 관한 것이다.'
)

doc.add_heading('3. 발명의 배경', 1)
doc.add_paragraph(
    'Amazon Alexa, Google Assistant, Apple Siri 등 기존 멀티기기 AI 시스템은 '
    '자사 에코시스템 내 동일 벤더 기기 간에만 세션 연속성을 제공하며, '
    '이기종 벤더의 AI 채널 간 페르소나 공유나 컨텍스트 자동 인계는 지원하지 않는다. '
    '또한 사용자가 직접 채널(기기)을 선택해야 하며, AI가 상황에 따라 최적 채널을 '
    '자동으로 선택하는 라우팅 기능은 존재하지 않는다.'
)

doc.add_heading('4. 발명의 목적', 1)
for i, g in enumerate([
    '이기종 AI 에이전트 채널 간 단일 AI 페르소나 유지',
    '채널 전환 시 사용자 컨텍스트 자동 인계 (세션 연속성)',
    'AI가 상황에 따라 최적 채널을 자동 선택하는 라우팅',
    '플랫폼 독립적 파일 시스템 기반 컨텍스트 동기화',
    '스마트폰 단말기를 HW 기반으로 한 발명 성립성 확보',
], 1):
    doc.add_paragraph(f'{i}. {g}')

doc.add_heading('5. 핵심 구성요소', 1)
doc.add_paragraph(
    '[채널 1] 스마트폰 모바일 앱 ← HW(단말기)+SW 결합\n'
    '  역할: 이동 중 원격 작업 지시, 질문\n'
    '       ↕\n'
    '[채널 2] 메신저 봇 (텔레그램 등)\n'
    '  역할: 알림, 빠른 조회, 모니터링\n'
    '       ↕\n'
    '[채널 3] 데스크탑 AI CLI\n'
    '  역할: 깊은 작업, 코드 작성, 파일 관리\n'
    '       ↕\n'
    '[공유 컨텍스트 레이어]\n'
    '  - 실시간 상태 동기화 (Firebase 등)\n'
    '  - 파일 시스템 기반 컨텍스트 인계\n'
    '  - AI 자동 채널 라우팅 모듈\n'
    '  - 단일 AI 페르소나 관리'
)

doc.add_heading('6. 신규성 및 차별점', 1)
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
h = table.rows[0].cells
h[0].text = '구분'; h[1].text = '기존 기술 (Amazon/Google 등)'; h[2].text = '본 발명'
for i, (a,b,c) in enumerate([
    ('채널 범위', '자사 에코시스템 전용', '이기종 벤더 채널 통합'),
    ('페르소나', '채널/기기마다 별개', '전 채널 단일 페르소나'),
    ('채널 선택', '사용자가 직접 선택', 'AI가 상황 판단 후 자동 라우팅'),
    ('컨텍스트 인계', '동일 벤더만 가능 (Apple Handoff 등)', '플랫폼 독립 파일 기반 자동 인계'),
], 1):
    r = table.rows[i].cells
    r[0].text = a; r[1].text = b; r[2].text = c

doc.add_paragraph()
doc.add_heading('6-1. 선행기술 공백 (사복이 조사 결과)', 2)
for item in [
    '이기종 AI 채널 간 단일 페르소나 유지 → 미발견',
    'AI 자동 채널 라우팅 (역할 기반) → 미발견',
    '파일 시스템 기반 플랫폼 독립 컨텍스트 동기화 → 미발견',
]:
    doc.add_paragraph(f'  ✅ {item}')

doc.add_heading('7. 청구항 초안', 1)
doc.add_paragraph().add_run('독립항').bold = True
doc.add_paragraph(
    '"모바일 단말기에 설치된 AI 에이전트 앱, 메신저 봇 서버, 데스크탑 AI 에이전트를 포함하는 '
    '복수의 이기종 AI 채널이 공유 컨텍스트 저장소를 통해 단일 사용자 컨텍스트와 AI 페르소나를 동기화하며, '
    'AI 라우팅 모듈이 사용자의 상황 및 요청 유형에 따라 최적 채널을 자동 선택하여 응답하되, '
    '채널 전환 시에도 이전 채널의 컨텍스트가 자동 인계되어 사용자 경험의 연속성이 유지되는 것을 '
    '특징으로 하는 퍼스널 AI 오케스트레이션 시스템"'
)
doc.add_paragraph()
for title, content in [
    ('종속항 1 — 자동 라우팅',
     '"제1항에 있어서, 라우팅 모듈은 요청의 복잡도, 사용자 위치, 기기 가용성을 분석하여 단순 조회는 메신저 봇, 이동 중 요청은 모바일 앱, 복잡한 작업은 데스크탑으로 자동 분기하는 것을 특징으로 하는 시스템"'),
    ('종속항 2 — IoT/완구 확장',
     '"제1항에 있어서, 복수의 채널은 IoT 기기 또는 AI 음성 완구를 추가로 포함하며, 해당 기기와도 동일한 페르소나 및 컨텍스트가 공유되는 것을 특징으로 하는 시스템"'),
]:
    doc.add_paragraph().add_run(title).bold = True
    doc.add_paragraph(content)

doc.add_heading('8. 특허 가능성', 1)
doc.add_paragraph(
    '중간~높음 — 스마트폰 단말기(HW)와 결합으로 발명 성립성 확보. '
    '이기종 AI 채널 통합 + 자동 라우팅 조합의 선행기술 미발견. '
    '11_#1(호두 AI)와 독립된 별개 특허로 출원 가능.'
)

doc.add_heading('9. 활용 방안', 1)
for m in [
    '개인 AI 비서 서비스 플랫폼',
    '스마트폰 제조사 번들 AI 시스템 (HW 바로보기님 직무 연관)',
    '기업용 멀티채널 AI 고객 응대 시스템',
    '스마트홈 + 모바일 통합 AI 허브',
]:
    doc.add_paragraph(f'• {m}')

doc.add_heading('10. 다음 단계', 1)
for s in [
    'KIPRIS 직접 선행검색',
    '변리사 상담 (11_#1 호두 AI와 별개 출원 전략 논의)',
    '국내 특허 출원 (KR)',
]:
    doc.add_paragraph(f'☐  {s}')

doc.add_paragraph()
doc.add_paragraph('* 본 문서는 특허 출원을 위한 기초 개념 정리용이며, 법적 효력은 없습니다.').italic = True

doc.save('D:/AI/11_특허아이디어/11_7_멀티채널AI_오케스트레이션_v0.2.docx')
print('11_#7 v0.2 완료')
