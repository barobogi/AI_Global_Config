# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 제목
title = doc.add_heading('호두 AI — 발명 기본 개념서', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.add_run('작성일: 2026-06-29  |  작성자: 이한복 (barobogi)  |  버전: v0.1 초안').italic = True
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# 1
doc.add_heading('1. 발명의 명칭', 1)
p = doc.add_paragraph()
p.add_run('가족 대화 누적 학습 기반 캐릭터 페르소나 AI 인형 시스템').bold = True
doc.add_paragraph('(영문: AI Character Doll System Based on Family Conversation Accumulative Learning)')

# 2
doc.add_heading('2. 기술 분야', 1)
doc.add_paragraph(
    'WiFi 통신 기능을 갖춘 무선 스피커 및 마이크 모듈을 캐릭터 인형 외형에 통합하고, '
    '가족 구성원 다수의 대화 데이터를 클라우드 AI가 장기 누적 학습하여 '
    '특정 캐릭터 페르소나로 응답하는 AI 인형 시스템에 관한 것이다.'
)

# 3
doc.add_heading('3. 발명의 배경', 1)
doc.add_paragraph(
    '기존 AI 스피커(Amazon Alexa, 카카오미니 등)는 불특정 다수를 대상으로 하는 범용 응답 시스템으로, '
    '특정 가족의 언어 패턴, 호칭, 내부 문화를 학습하지 못한다. '
    '또한 기존 AI 완구는 단일 사용자(주로 아동 1인) 대상 개인화에 그치며, '
    '가족 전체의 상호작용 데이터를 통합 학습하는 기술은 상용화된 사례가 없다.'
)
doc.add_paragraph(
    '가정 내에서 캐릭터 인형을 인격화하여 대화하는 문화(인형에게 말 걸기, 인형 목소리 흉내 등)가 '
    '존재하나, 이를 AI로 실제 구현한 제품은 없다.'
)

# 4
doc.add_heading('4. 발명의 목적', 1)
goals = [
    '가족 구성원 복수 화자의 대화 데이터를 장기 누적 학습',
    '특정 캐릭터의 페르소나를 일관되게 유지하며 응답',
    '대화가 쌓일수록 해당 가족에 특화된 응답을 생성하여 관계 발전을 표현',
    'WiFi 내장 스피커/마이크 모듈을 인형 외형에 통합한 독립 기기 구현',
]
for i, g in enumerate(goals, 1):
    doc.add_paragraph(f'{i}. {g}')

# 5
doc.add_heading('5. 핵심 구성요소', 1)
doc.add_paragraph(
    '[캐릭터 인형 외형]\n'
    '       ↕\n'
    '[WiFi 내장 스피커 + 마이크 모듈]\n'
    '       ↕  (WiFi 통신)\n'
    '[클라우드 AI 서버]\n'
    '  - STT (음성 → 텍스트)\n'
    '  - 가족 대화 데이터 누적 저장\n'
    '  - 캐릭터 페르소나 관리 모듈\n'
    '  - 가족 특화 응답 생성 (LLM)\n'
    '  - TTS (텍스트 → 음성)\n'
    '       ↕\n'
    '[응답 음성 출력 → 스피커]'
)

# 6
doc.add_heading('6. 신규성 및 차별점', 1)
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '구분'
hdr[1].text = '기존 기술'
hdr[2].text = '본 발명'
rows_data = [
    ('학습 대상', '단일 사용자 또는 아동 1인', '가족 구성원 다수 (복수 화자)'),
    ('학습 기간', '세션 기반 단기', '장기 누적 (관계 발전 표현)'),
    ('캐릭터성', '범용 AI 어시스턴트', '특정 캐릭터 페르소나 고정 유지'),
    ('개인화 단위', '개인 프로필', '가족 단위 언어·문화 특화'),
    ('형태', '스피커 또는 로봇', '캐릭터 인형 외형 통합'),
]
for i, (a, b, c) in enumerate(rows_data, 1):
    r = table.rows[i].cells
    r[0].text = a
    r[1].text = b
    r[2].text = c

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('핵심 청구항 후보: ').bold = True
p.add_run(
    '복수의 가족 구성원 음성 데이터를 장기 누적 학습하여 '
    '해당 가족에 특화된 캐릭터 페르소나 응답을 생성하는 '
    'AI 인형 시스템 및 그 방법'
)

# 7
doc.add_heading('7. 활용 방안 및 시장성', 1)
markets = [
    '타겟: 어린 자녀가 있는 가정, 인형을 키우는 1~2인 가구',
    '캐릭터 IP 연계: 카카오프렌즈, 라인프렌즈 등 기존 IP와 라이선스 결합 가능',
    '구독 모델: 클라우드 AI 학습 서비스 월정액',
    '확장: 반려동물 AI 인형, 조부모-손자녀 연결 AI 인형 등',
]
for m in markets:
    doc.add_paragraph(f'• {m}')

# 8
doc.add_heading('8. 다음 단계', 1)
steps = [
    '변리사 상담 (청구항 범위 확정)',
    'KIPRIS 추가 선행 검색 (한국어 키워드 정밀 검색)',
    '프로토타입 구현 (라즈베리파이 + Claude API)',
    '국내 특허 출원 (KR)',
]
for s in steps:
    doc.add_paragraph(f'☐  {s}')

doc.add_paragraph()
p = doc.add_paragraph('* 본 문서는 특허 출원을 위한 기초 개념 정리용이며, 법적 효력은 없습니다.')
p.italic = True

doc.save('D:/AI/DEV_ITEM/호두AI_발명개념서_v0.1.docx')
print('Word 파일 생성 완료')
