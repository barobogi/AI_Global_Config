# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 제목
t = doc.add_heading('호두 AI — 발명 핵심 요약 (1 Page)', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = doc.add_paragraph()
meta.add_run('작성일: 2026-06-29  |  작성자: 이한복 (barobogi)  |  버전: v0.31 (요약본)').italic = True
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# 발명명
doc.add_heading('발명명', 2)
doc.add_paragraph().add_run('가족 대화 누적 학습 기반 캐릭터 페르소나 AI 인형 시스템').bold = True

# 한 줄 요약
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('한 줄 요약: ').bold = True
p.add_run('"우리 가족과 대화할수록 친해지는 AI 인형"')

# 핵심 차별점 3가지
doc.add_heading('핵심 차별점 (선행기술 미발견)', 2)
for d in [
    '① 가족 복수 화자 누적 학습 (기존: 단일 사용자)',
    '② 누적 대화량 → 친밀도 L1~L5 단계적 강화 (무뚝뚝→애착)',
    '③ 물리 인형 + 모바일 앱 실시간 동기화',
]:
    doc.add_paragraph(d)

# 친밀도 수식
doc.add_heading('친밀도 산출 수식', 2)
doc.add_paragraph('I = 0.4C + 0.3S + 0.2T + 0.1R')
doc.add_paragraph('(C:대화횟수  S:감정점수  T:주제다양성  R:최근성)')

# 독립항
doc.add_heading('핵심 청구항 (독립항)', 2)
doc.add_paragraph(
    '"복수 가족 구성원 음성 데이터를 장기 누적 학습하여, '
    '특정 캐릭터 페르소나로 응답하되, '
    '누적 대화량에 따라 친밀도 표현이 단계적으로 강화되는 AI 인형 시스템"'
)

# 선행기술 vs 본 발명
doc.add_heading('선행기술 대비 우위', 2)
table = doc.add_table(rows=4, cols=3)
table.style = 'Table Grid'
h = table.rows[0].cells
h[0].text = '경쟁사'; h[1].text = '한계'; h[2].text = '본 발명 우위'
for i, (a,b,c) in enumerate([
    ('Sony LOVOT', '단일 사용자, 친밀도 없음', '가족 전체 + L1~L5'),
    ('Replika', '앱만, 물리기기 없음', '물리 인형+모바일 하이브리드'),
    ('Amazon Echo', '음성인식만, 감정없음', '페르소나+친밀도 완전 결합'),
], 1):
    r = table.rows[i].cells
    r[0].text = a; r[1].text = b; r[2].text = c

# 다음 단계
doc.add_paragraph()
doc.add_heading('다음 단계', 2)
for s in ['① KIPRIS 선행검색', '② 변리사 상담 (최우선)', '③ 국내 출원 (KR)']:
    doc.add_paragraph(f'  {s}')

doc.add_paragraph()
doc.add_paragraph('* 상세 내용: 호두AI_발명개념서_v0.3.docx 참조  |  법적 효력 없음').italic = True

doc.save('D:/AI/11_특허아이디어/호두AI_발명개념서_v0.31_요약.docx')
print('v0.31 완료')
