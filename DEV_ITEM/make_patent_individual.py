# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def make_doc(filename, title_kr, title_en, field, background, purpose_list, structure, diff_rows, claim, possibility, market):
    doc = Document()
    t = doc.add_heading(title_kr, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph()
    meta.add_run('작성일: 2026-06-29  |  작성자: 이한복 (barobogi)  |  버전: v0.1 초안').italic = True
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    doc.add_heading('1. 발명의 명칭', 1)
    doc.add_paragraph().add_run(title_kr).bold = True
    doc.add_paragraph(f'(영문: {title_en})')

    doc.add_heading('2. 기술 분야', 1)
    doc.add_paragraph(field)

    doc.add_heading('3. 발명의 배경', 1)
    doc.add_paragraph(background)

    doc.add_heading('4. 발명의 목적', 1)
    for i, p in enumerate(purpose_list, 1):
        doc.add_paragraph(f'{i}. {p}')

    doc.add_heading('5. 핵심 구성요소', 1)
    doc.add_paragraph(structure)

    doc.add_heading('6. 신규성 및 차별점', 1)
    table = doc.add_table(rows=len(diff_rows)+1, cols=3)
    table.style = 'Table Grid'
    h = table.rows[0].cells
    h[0].text = '구분'; h[1].text = '기존 기술'; h[2].text = '본 발명'
    for i, (a, b, c) in enumerate(diff_rows, 1):
        r = table.rows[i].cells
        r[0].text = a; r[1].text = b; r[2].text = c
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('핵심 청구항 후보: ').bold = True
    p.add_run(f'"{claim}"')

    doc.add_heading('7. 특허 가능성', 1)
    doc.add_paragraph(possibility)

    doc.add_heading('8. 활용 방안', 1)
    for m in market:
        doc.add_paragraph(f'• {m}')

    doc.add_paragraph()
    doc.add_paragraph('* 본 문서는 특허 출원을 위한 기초 개념 정리용이며, 법적 효력은 없습니다.').italic = True

    doc.save(f'D:/AI/DEV_ITEM/{filename}')
    print(f'{filename} 완료')


# ① AI 세션 간 자동 컨텍스트 동기화
make_doc(
    '특허①_AI세션_컨텍스트_동기화_v0.1.docx',
    '① AI 에이전트 세션 간 자동 컨텍스트 동기화 시스템',
    'Automatic Context Synchronization System Between AI Agent Sessions',
    'AI 어시스턴트 세션이 종료되거나 전환될 때 파일 시스템 감시를 통해 다음 AI 에이전트가 판독 가능한 컨텍스트 파일을 자동 생성·동기화하는 시스템에 관한 것.',
    '현재 AI 어시스턴트(Claude, GPT 등)는 세션 종료 시 컨텍스트가 소멸되며, 다음 세션 또는 다른 AI 에이전트가 이전 작업을 이어받으려면 사용자가 수동으로 내용을 전달해야 한다. 이로 인해 AI 협업 시 반복적인 맥락 재설명이 필요하고 작업 연속성이 떨어진다.',
    ['AI 세션 종료/전환 시 컨텍스트 자동 보존', '복수의 이기종 AI 에이전트 간 표준화된 컨텍스트 전달', '사용자 개입 없는 완전 자동화된 세션 브릿지 구현'],
    '파일 시스템 감시 모듈 (watchdog)\n  ↓ 변경 감지\n컨텍스트 추출 모듈 (현재 세션 요약)\n  ↓\n구조화된 브릿지 파일 자동 생성 (Markdown/JSON)\n  ↓\n다음 AI 에이전트 세션 시작 시 자동 로드',
    [
        ('컨텍스트 보존', '세션 종료 시 소멸', '파일로 자동 영속화'),
        ('전달 방식', '사용자 수동 복사붙여넣기', '감시 기반 자동 생성'),
        ('대상', '단일 AI 플랫폼 내부', '이기종 AI 에이전트 간'),
        ('플랫폼', '특정 서비스 종속', '파일 시스템 기반 독립'),
    ],
    '파일 시스템 감시 모듈이 AI 세션 종료 이벤트를 감지하여 다음 AI 에이전트 세션이 판독 가능한 구조화된 컨텍스트 파일을 자동 생성하는 방법 및 시스템',
    '중간 — 순수 SW 발명으로 단독 특허는 어려우나, 파일 감시+자동생성+AI 에이전트 연동을 결합한 시스템 청구항으로 가능성 있음. AI 협업 도구 특허 출원 증가 추세.',
    ['AI 개발 도구 플랫폼 내장 기능', '기업용 AI 에이전트 오케스트레이션 시스템', 'IDE/코드에디터 플러그인'],
)

# ② 외부 AI 위임 실행 패턴
make_doc(
    '특허②_외부AI_위임실행_v0.1.docx',
    '② 비용 등급 기반 AI 역할 분리 위임 실행 시스템',
    'Cost-tier Based AI Role Separation and Delegation Execution System',
    '고비용 고성능 AI 모델이 작업을 설계·판단하고, 저비용 AI 모델이 반복 실행을 담당하며, 결과를 파일로 교환하는 다계층 AI 실행 위임 아키텍처에 관한 것.',
    '단일 고성능 AI 모델로 반복 작업을 처리하면 불필요한 비용이 발생한다. 기존 AI 오케스트레이션 도구는 역할 분리 시 비용 효율성을 고려하지 않으며, 모델 간 인터페이스가 복잡하다.',
    ['AI 작업의 설계/판단과 반복 실행을 명시적으로 분리', '비용 효율적인 다계층 AI 실행 파이프라인 구현', '파일 기반 표준 인터페이스로 플랫폼 독립성 확보'],
    '주 AI 모델 (고비용, 설계+판단)\n  ↓ 작업 명세 파일 생성\n보조 AI 모델 (저비용, 반복 실행)\n  ↓ 결과 파일 저장\n주 AI 모델 (결과 검증 + 최종 판단)\n  ↓\n다음 작업 명세 생성 (반복)',
    [
        ('역할 분리', '단일 모델 처리', '비용 등급별 명시적 분리'),
        ('인터페이스', 'API 직접 호출', '파일 시스템 기반 표준화'),
        ('비용 최적화', '고려 없음', '저비용 모델에 반복 위임'),
        ('품질 보증', '단일 검증', '2단계 검증 (실행+판단 분리)'),
    ],
    '제1 AI 모델이 작업 명세를 생성하고, 비용이 낮은 제2 AI 모델이 반복 실행 후 결과 파일을 저장하며, 제1 AI 모델이 결과를 검증하는 비용 등급 기반 다단계 AI 실행 위임 방법',
    '중간 — 오케스트레이션 개념 자체는 선행 기술 있으나, 비용 등급 기반 역할 분리+파일 인터페이스 조합은 신규성 주장 가능. 생성형 AI 확산으로 비용 최적화 특허 수요 증가 중.',
    ['엔터프라이즈 AI 워크플로우 플랫폼', 'LLM 오케스트레이션 미들웨어', '클라우드 AI 서비스 비용 최적화 도구'],
)

# ③ 시장 기반 자동학습
make_doc(
    '특허③_시장기반_신호자동학습_v0.1.docx',
    '③ 인간 판단 배제 시장 결과 기반 투자 신호 자동 최적화 시스템',
    'Human Bias-Free Automated Investment Signal Optimization System Using Market Outcome Feedback',
    '투자 신호 생성 후 실제 시장 가격 변동을 자동 수집하여, 인간 판단 개입 없이 신호 패턴별 가중치를 자동 조정하는 투자 신호 최적화 시스템에 관한 것.',
    '기존 알고리즘 트레이딩 최적화는 강화학습(대규모 인프라 필요), 전문가 피드백(인간 편향 개입), 백테스팅(사후 검증만 가능) 방식이 주를 이룬다. 특히 인간 피드백 기반 학습은 투자자 개인의 심리적 편향이 시스템에 학습되는 문제가 있다.',
    ['시장 가격만을 유일한 학습 신호로 사용하여 인간 편향 완전 배제', '개인 투자자 수준에서 구현 가능한 경량 자동학습 파이프라인 구현', '단순 명확한 피드백 기준(N일 후 수익률)으로 복잡한 보상 함수 불필요'],
    '투자 신호 생성 모듈 (AI 등급 산출)\n  ↓ 신호 + 날짜 저장\n시장 데이터 자동 수집 (N일 후)\n  ↓\n패턴별 실제 수익률 계산\n  ↓\n신호 가중치 자동 조정 (인간 개입 없음)\n  ↓\n다음 신호 생성에 반영',
    [
        ('학습 신호', '전문가 판단 또는 복잡한 보상함수', '시장 가격(객관적 데이터)만 사용'),
        ('인간 편향', '피드백 과정에서 개입', '완전 배제'),
        ('인프라', '대규모 RL 시스템 필요', '경량 파이프라인 (개인 운용 가능)'),
        ('적용 대상', '기관 투자자', '개인 투자자'),
    ],
    '투자 신호 생성 후 일정 기간 경과 시 시장 가격 데이터를 자동 수집하여 인간 판단 개입 없이 신호 패턴별 가중치를 갱신하는 자동화된 투자 신호 최적화 방법 및 시스템',
    '중간~높음 — 금융+SW 결합으로 BM특허 가능. 인간 편향 배제라는 구체적 기술 효과가 신규성 주장에 유리. 개인 투자자 대상 AI 투자 도구 특허 출원 활발.',
    ['개인 투자자용 AI 투자 보조 서비스', '증권사 로보어드바이저 고도화', '핀테크 스타트업 핵심 기술'],
)

# ④ 4계층 아키텍처
make_doc(
    '특허④_4계층_AI협업_아키텍처_v0.1.docx',
    '④ 인간-AI 협업을 위한 4계층 지속성 컨텍스트 아키텍처',
    'Four-Layer Persistent Context Architecture for Human-AI Collaboration',
    '인간과 AI가 장기간 협업할 때 AI 세션 단절 문제를 해결하기 위해, 전역 규칙·기술 이력·재사용 코드·프로젝트 현황의 4계층으로 구성된 파일 시스템 기반 컨텍스트 관리 아키텍처에 관한 것.',
    'AI 어시스턴트는 세션마다 초기화되어 이전 작업 맥락을 잃는다. 기존 해결책(단일 규칙 파일, RAG 등)은 계층 구조 없이 모든 정보를 혼재시켜 관리 효율이 낮고, 재사용 가능한 지식의 체계적 축적이 어렵다.',
    ['AI 세션 단절 문제를 4계층 파일 구조로 체계적 해결', '계층 간 승격(promotion) 메커니즘으로 지식 자산 누적', '다수의 AI 에이전트가 동일 컨텍스트를 공유하는 협업 구조 실현'],
    'Layer 0: 전역 규칙 파일 (모든 프로젝트 공통 적용)\n  ↓\nLayer 1: 기술 이력 저장소 (프로젝트별 개발 기술 기록)\n  ↓ (검증 후 승격)\nLayer 2: 재사용 코드 창고 (검증된 모듈 보관)\n  ↓\nLayer 3: 프로젝트 현황 파일 (세션 재시작 시 복원)',
    [
        ('계층 구조', '없음 (단일 파일)', '명시적 4계층 역할 분리'),
        ('지식 축적', '세션마다 소멸', '계층 간 승격으로 영속화'),
        ('재사용', '수동 복사', '레이어2 창고에서 자동 제안'),
        ('적용 범위', '단일 프로젝트', '전체 프로젝트 공통 적용'),
    ],
    '전역 규칙, 기술 이력, 재사용 코드, 프로젝트 현황의 4계층으로 구성되며 계층 간 승격 메커니즘을 포함하는 파일 시스템 기반 인간-AI 협업 컨텍스트 관리 방법 및 시스템',
    '낮음~중간 — 추상적 아키텍처에 가까워 특허보다 방어적 공개(defensive publication)가 적합할 수 있음. 단, 파일 구조+자동화 도구+승격 메커니즘을 구체화하면 청구 가능성 있음.',
    ['AI 개발 환경 통합 도구 (IDE 플러그인)', '기업용 AI 에이전트 관리 플랫폼', 'AI 협업 프레임워크'],
)
