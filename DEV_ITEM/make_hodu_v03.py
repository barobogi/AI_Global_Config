# -*- coding: utf-8 -*-
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 제목
t = doc.add_heading('호두 AI — 발명 기본 개념서', 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = doc.add_paragraph()
meta.add_run('작성일: 2026-06-29  |  작성자: 이한복 (barobogi)  |  버전: v0.3').italic = True
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 버전 이력
doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('버전 이력: ').bold = True
p.add_run('v0.1 초안 → v0.2 핵심 청구항 3가지 통합 → v0.3 선행기술 조사 + 청구항 구체화')
doc.add_paragraph()

# 1. 발명의 명칭
doc.add_heading('1. 발명의 명칭', 1)
doc.add_paragraph().add_run('가족 대화 누적 학습 기반 캐릭터 페르소나 AI 인형 시스템').bold = True
doc.add_paragraph('(영문: AI Character Doll System Based on Family Conversation Accumulative Learning with Progressive Intimacy Expression)')

# 2. 기술 분야
doc.add_heading('2. 기술 분야', 1)
doc.add_paragraph(
    'WiFi 통신 기능을 갖춘 무선 스피커 및 마이크 모듈을 캐릭터 인형 외형에 통합하고, '
    '가족 구성원 다수의 대화 데이터를 클라우드 AI가 장기 누적 학습하여 '
    '특정 캐릭터 페르소나로 응답하며, 대화 누적량에 따라 친밀도 표현이 단계적으로 강화되는 '
    'AI 인형 시스템에 관한 것이다.'
)

# 3. 발명의 배경
doc.add_heading('3. 발명의 배경', 1)
doc.add_paragraph(
    '기존 AI 스피커(Amazon Alexa, 카카오미니 등)는 불특정 다수를 대상으로 하는 범용 응답 시스템으로, '
    '특정 가족의 언어 패턴, 호칭, 내부 문화를 학습하지 못한다. '
    '또한 기존 AI 완구는 단일 사용자(주로 아동 1인) 대상 개인화에 그치며, '
    '가족 전체의 상호작용 데이터를 통합 학습하는 기술은 상용화된 사례가 없다.'
)
doc.add_paragraph(
    '특히 기존 AI 시스템은 대화가 쌓여도 응답 방식이 변하지 않아 관계가 깊어지는 느낌을 표현하지 못한다. '
    '처음에는 무뚝뚝하다가 대화가 쌓일수록 친근해지는 친밀도 표현 강화 기능은 구현된 사례가 없다.'
)

# 4. 발명의 목적
doc.add_heading('4. 발명의 목적', 1)
for i, g in enumerate([
    '가족 구성원 복수 화자의 대화 데이터를 장기 누적 학습',
    '특정 캐릭터의 페르소나를 일관되게 유지하며 응답',
    '누적 대화량에 따라 친밀도 표현이 단계적으로 강화 (L1 무뚝뚝 → L5 애착)',
    'WiFi 내장 스피커/마이크 모듈을 인형 외형에 통합한 독립 기기 구현',
    '모바일 앱과 연동하여 기기 무관 동일 AI 페르소나 연속 유지 (v0.3 추가)',
], 1):
    doc.add_paragraph(f'{i}. {g}')

# 5. 핵심 구성요소
doc.add_heading('5. 핵심 구성요소', 1)
doc.add_paragraph(
    '[캐릭터 인형 외형]\n'
    '       ↕\n'
    '[WiFi 내장 스피커 + 마이크 모듈]\n'
    '       ↕  (WiFi 통신)\n'
    '[클라우드 AI 서버]\n'
    '  - STT (음성 → 텍스트)\n'
    '  - 화자 인식 + 가족 관계 그래프 관리\n'
    '  - 가족 대화 데이터 누적 저장\n'
    '  - 친밀도 레벨 관리 모듈 (I = 0.4C + 0.3S + 0.2T + 0.1R)\n'
    '  - 캐릭터 페르소나 관리 (고정 코어 + 가변 톤 계층 분리)\n'
    '  - 가족 특화 응답 생성 (LLM)\n'
    '  - TTS (텍스트 → 음성)\n'
    '       ↕  (실시간 동기화)\n'
    '[모바일 앱] ← v0.3 추가'
)

# 6. 신규성 및 차별점
doc.add_heading('6. 신규성 및 차별점', 1)
table = doc.add_table(rows=6, cols=3)
table.style = 'Table Grid'
h = table.rows[0].cells
h[0].text = '구분'; h[1].text = '기존 기술'; h[2].text = '본 발명'
for i, (a,b,c) in enumerate([
    ('학습 대상', '단일 사용자 또는 아동 1인', '가족 구성원 다수 (복수 화자)'),
    ('친밀도 표현', '고정된 응답 방식', '누적 대화량에 따라 L1~L5 단계적 강화'),
    ('캐릭터성', '범용 AI 어시스턴트', '특정 캐릭터 페르소나 고정 유지'),
    ('기기 연동', '단일 기기', '물리 인형 + 모바일 앱 실시간 동기화'),
    ('형태', '스피커 또는 로봇', '캐릭터 인형 외형 통합'),
], 1):
    r = table.rows[i].cells
    r[0].text = a; r[1].text = b; r[2].text = c

# 친밀도 레벨 표
doc.add_paragraph()
doc.add_heading('6-1. 친밀도 레벨 산출 함수', 2)
doc.add_paragraph('I(u,t) = 0.4 × C_u + 0.3 × S_u + 0.2 × T_u + 0.1 × R_u')
doc.add_paragraph('C: 누적 대화 횟수  |  S: 평균 감정 점수  |  T: 주제 다양성  |  R: 최근성 가중치')

level_table = doc.add_table(rows=6, cols=3)
level_table.style = 'Table Grid'
lh = level_table.rows[0].cells
lh[0].text = '레벨'; lh[1].text = '점수 범위'; lh[2].text = '응답 톤'
for i, (a,b,c) in enumerate([
    ('L1', '0.0~0.2', '정중한 무뚝뚝 — "네, 알겠습니다."'),
    ('L2', '0.2~0.4', '친절 — "그렇군요~"'),
    ('L3', '0.4~0.6', '친근 — "오~ 재밌다!"'),
    ('L4', '0.6~0.8', '애교 — "정말요?? 좋아요~!"'),
    ('L5', '0.8~1.0', '애착 — "보고 싶었어요!"'),
], 1):
    r = level_table.rows[i].cells
    r[0].text = a; r[1].text = b; r[2].text = c

# 선행기술 조사
doc.add_paragraph()
doc.add_heading('6-2. 선행기술 조사 결과 (v0.3 신규)', 2)
doc.add_paragraph().add_run('선행기술 공백 — 아래 조합은 미발견:').bold = True
for item in [
    '복수 화자 × 개별 친밀도 × 가족 관계 그래프',
    '대화 횟수+감정+주제다양성 가중합 → 레벨 함수',
    '물리 완구 + 모바일 앱 간 친밀도·페르소나 실시간 동기화',
]:
    doc.add_paragraph(f'  ✅ {item}')

doc.add_paragraph()
doc.add_paragraph().add_run('⚠️ 유사 선행기술 (회피 가능):').bold = True
st = doc.add_table(rows=5, cols=3)
st.style = 'Table Grid'
sh = st.rows[0].cells
sh[0].text = '출원인'; sh[1].text = '핵심 내용'; sh[2].text = '회피 포인트'
for i, (a,b,c) in enumerate([
    ('Embodied Inc.', 'AI 로봇 감정 적응', '단일 사용자, 친밀도 함수 없음'),
    ('Sony', '반려 로봇 감정 추적', '가족 단위 미지원, LLM 미사용'),
    ('Replika', 'AI 관계 진전 시스템', '물리 기기 없음, 단일 사용자'),
    ('Amazon', '복수 화자 음성 프로필', '친밀도 개념 없음, 완구 아님'),
], 1):
    r = st.rows[i].cells
    r[0].text = a; r[1].text = b; r[2].text = c

# 7. 청구항 초안
doc.add_heading('7. 청구항 초안 (v0.3 신규)', 1)

doc.add_paragraph().add_run('독립항').bold = True
doc.add_paragraph(
    '"음성 입출력 모듈, 화자 인식 모듈, 클라우드 언어모델 연동 모듈, 복수 화자별 친밀도 데이터베이스를 포함하는 '
    'AI 대화형 완구 시스템으로서, 각 화자의 누적 대화 패턴을 분석하여 친밀도 레벨을 산출하고, '
    '해당 레벨에 대응하는 응답 톤 파라미터를 언어모델의 시스템 프롬프트에 주입함으로써 '
    '응답 어조를 단계적으로 변화시키는 것을 특징으로 하는 시스템"'
)

for n, title, content in [
    (1, '종속항 1 — 친밀도 알고리즘',
     '"제1항에 있어서, 친밀도 레벨은 대화 횟수(C), 감정 분석 점수(S), 주제 다양성 지수(T), 최근성 가중치(R)의 가중합 I = 0.4C + 0.3S + 0.2T + 0.1R 로 산출되며, L1~L5의 레벨로 단계적 상승하는 것을 특징으로 하는 시스템"'),
    (2, '종속항 2 — 가족 관계 그래프',
     '"제1항에 있어서, 복수의 화자 프로필을 가족 관계 그래프로 연결하여 저장하고, 각 화자 간 호칭 체계 및 가족 내부 에피소드를 공유 맥락으로 추출·저장하여 응답 생성 시 활용하는 것을 특징으로 하는 시스템"'),
    (3, '종속항 3 — 페르소나 일관성',
     '"제1항에 있어서, 완구의 핵심 정체성을 정의하는 고정 코어 레이어와 친밀도 레벨에 따라 변화하는 가변 톤 레이어를 계층적으로 분리하는 것을 특징으로 하는 시스템"'),
    (4, '종속항 4 — 물리+모바일 하이브리드',
     '"제1항에 있어서, 모바일 애플리케이션과 연동하여 물리 완구 기기와 모바일 기기 사이에서 화자별 친밀도 레벨, 가족 관계 그래프, 페르소나 상태를 실시간 동기화함으로써 기기 이탈 상황에서도 동일 AI 페르소나와의 대화를 연속 유지하는 것을 특징으로 하는 시스템"'),
]:
    doc.add_paragraph()
    doc.add_paragraph().add_run(title).bold = True
    doc.add_paragraph(content)

# 8. 활용 방안
doc.add_heading('8. 활용 방안 및 시장성', 1)
for m in [
    '타겟: 어린 자녀가 있는 가정, 캐릭터 인형에 애착이 있는 가족',
    '캐릭터 IP 연계: 카카오프렌즈, 라인프렌즈 등 기존 IP 라이선스 결합',
    '구독 모델: 클라우드 AI 학습 서비스 월정액',
    '확장: 반려동물 AI 인형, 조부모-손자녀 연결 AI 인형',
]:
    doc.add_paragraph(f'• {m}')

# 9. 다음 단계
doc.add_heading('9. 다음 단계', 1)
for s in [
    'KIPRIS 직접 선행검색 (키워드: AI완구 복수화자, 감성로봇 친밀도 단계)',
    '변리사 상담 (청구항 범위 확정) ← 최우선',
    '프로토타입 구현 (라즈베리파이 + Claude API)',
    '국내 특허 출원 (KR) — 프로토타입 공개 전 반드시 먼저',
]:
    doc.add_paragraph(f'☐  {s}')

doc.add_paragraph()
doc.add_paragraph('* 본 문서는 특허 출원을 위한 기초 개념 정리용이며, 법적 효력은 없습니다.').italic = True

doc.save('D:/AI/11_특허아이디어/호두AI_발명개념서_v0.3.docx')
print('v0.3 완료')
